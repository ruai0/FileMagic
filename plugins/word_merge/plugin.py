import os
from typing import Dict, Any, List, Optional
from docx import Document
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QGroupBox, QMessageBox,
    QTableWidget, QTableWidgetItem, QHeaderView
)
from PyQt6.QtCore import Qt
from ui.widgets.file_drop_zone import FileDropZone
from ui.widgets.log_panel import LogPanel
from core.plugin_manager import BasePlugin
from ui.dialogs.file_dialog import FileDialog
from ui.dialogs.progress_dialog import ProgressDialog


class Plugin(BasePlugin):
    """Word合并插件"""
    
    def __init__(self):
        super().__init__()
        self.name = "Word合并"
        self.description = "多个Word文档合并为一个"
        self.version = "1.0.0"
        self.author = "ExcelTools"
        self._widget = None
        self._context: Optional[Dict[str, Any]] = None
        self.files: List[str] = []
    
    def initialize(self, context: Dict[str, Any]) -> None:
        super().initialize(context)
        self._context = context
    
    def get_widget(self):
        if self._widget is None:
            self._widget = self._create_widget()
        return self._widget
    
    def _create_widget(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        file_group = QGroupBox("Word文件列表")
        file_layout = QVBoxLayout()
        
        self.drop_zone = FileDropZone()
        self.drop_zone.files_dropped.connect(self._on_files_dropped)
        file_layout.addWidget(self.drop_zone)
        
        self.file_table = QTableWidget()
        self.file_table.setColumnCount(2)
        self.file_table.setHorizontalHeaderLabels(["文件名", "段落数"])
        self.file_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        file_layout.addWidget(self.file_table)
        
        btn_row = QHBoxLayout()
        add_btn = QPushButton("添加文件")
        add_btn.clicked.connect(self._on_add)
        btn_row.addWidget(add_btn)
        remove_btn = QPushButton("移除选中")
        remove_btn.clicked.connect(self._on_remove)
        btn_row.addWidget(remove_btn)
        up_btn = QPushButton("上移")
        up_btn.clicked.connect(self._on_up)
        btn_row.addWidget(up_btn)
        down_btn = QPushButton("下移")
        down_btn.clicked.connect(self._on_down)
        btn_row.addWidget(down_btn)
        clear_btn = QPushButton("清空")
        clear_btn.clicked.connect(self._on_clear)
        btn_row.addWidget(clear_btn)
        file_layout.addLayout(btn_row)
        
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        merge_btn = QPushButton("开始合并")
        merge_btn.setStyleSheet("background-color: #0078d4; color: white; padding: 10px; font-size: 14px;")
        merge_btn.clicked.connect(self._on_merge)
        layout.addWidget(merge_btn)
        
        self.log_panel = LogPanel()
        layout.addWidget(self.log_panel)
        
        return widget
    
    def _on_files_dropped(self, filepaths):
        for fp in filepaths:
            if fp.lower().endswith((".docx",)) and fp not in self.files:
                self.files.append(fp)
        self._update_table()
    
    def _on_add(self):
        filepaths = FileDialog.open_files(self._widget, "选择Word文件", "Word文件 (*.docx)")
        if filepaths:
            for fp in filepaths:
                if fp not in self.files:
                    self.files.append(fp)
            self._update_table()
    
    def _update_table(self):
        self.file_table.setRowCount(len(self.files))
        for row, fp in enumerate(self.files):
            self.file_table.setItem(row, 0, QTableWidgetItem(os.path.basename(fp)))
            try:
                doc = Document(fp)
                self.file_table.setItem(row, 1, QTableWidgetItem(str(len(doc.paragraphs))))
            except:
                self.file_table.setItem(row, 1, QTableWidgetItem("N/A"))
    
    def _on_remove(self):
        row = self.file_table.currentRow()
        if row >= 0 and row < len(self.files):
            self.files.pop(row)
            self._update_table()
    
    def _on_up(self):
        row = self.file_table.currentRow()
        if row > 0:
            self.files[row], self.files[row-1] = self.files[row-1], self.files[row]
            self._update_table()
            self.file_table.setCurrentCell(row-1, 0)
    
    def _on_down(self):
        row = self.file_table.currentRow()
        if row < len(self.files) - 1:
            self.files[row], self.files[row+1] = self.files[row+1], self.files[row]
            self._update_table()
            self.file_table.setCurrentCell(row+1, 0)
    
    def _on_clear(self):
        self.files.clear()
        self.file_table.setRowCount(0)
    
    def _on_merge(self):
        if len(self.files) < 2:
            QMessageBox.warning(self._widget, "警告", "请至少添加2个Word文件")
            return
        
        output_path = FileDialog.save_file(self._widget, "保存合并结果", "Word文件 (*.docx)", ".docx")
        if not output_path:
            return
        
        progress = ProgressDialog("合并中...")
        progress.show()
        
        try:
            merged_doc = Document()
            total = len(self.files)
            
            for idx, fp in enumerate(self.files):
                progress.set_status(f"合并 {idx+1}/{total}: {os.path.basename(fp)}")
                progress.set_progress(int((idx / total) * 100))
                
                src_doc = Document(fp)
                
                for para in src_doc.paragraphs:
                    merged_doc.add_paragraph(para.text, para.style)
                
                for table in src_doc.tables:
                    new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
                
                if idx < total - 1:
                    merged_doc.add_page_break()
                
                progress.add_log(f"已合并: {os.path.basename(fp)}")
            
            merged_doc.save(output_path)
            progress.complete("合并完成!")
            self.log_panel.info(f"已合并 {total} 个文件到: {output_path}")
            QMessageBox.information(self._widget, "成功", f"已合并 {total} 个Word文件")
        except Exception as e:
            progress.close()
            self.log_panel.error(f"合并失败: {str(e)}")
            QMessageBox.critical(self._widget, "错误", f"合并失败:\n{str(e)}")
