import os
import pandas as pd
from docx import Document
from typing import Dict, Any, Optional
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QComboBox, QGroupBox, QMessageBox,
    QLineEdit, QTextEdit
)
from PyQt6.QtCore import Qt
from ui.widgets.file_drop_zone import FileDropZone
from ui.widgets.log_panel import LogPanel
from core.plugin_manager import BasePlugin
from ui.dialogs.file_dialog import FileDialog
from ui.dialogs.progress_dialog import ProgressDialog


class Plugin(BasePlugin):
    """Word模板填充插件"""
    
    def __init__(self):
        super().__init__()
        self.name = "Word模板填充"
        self.description = "用Excel数据批量填充Word模板生成文档"
        self.version = "1.0.0"
        self.author = "ExcelTools"
        self._widget = None
        self._context = None
        self.template_file = None
        self.excel_file = None
        self.excel_df = None
    
    def initialize(self, context):
        super().initialize(context)
        self._context = context
    
    def get_widget(self):
        if self._widget is None:
            self._widget = self._create_widget()
        return self._widget
    
    def _create_widget(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        template_group = QGroupBox("Word模板文件")
        template_layout = QVBoxLayout()
        self.template_drop = FileDropZone()
        self.template_drop.files_dropped.connect(self._on_template_dropped)
        self.template_drop.setMaximumHeight(80)
        template_layout.addWidget(self.template_drop)
        self.template_label = QLabel("未选择模板")
        self.template_label.setStyleSheet("padding: 5px; color: #666;")
        template_layout.addWidget(self.template_label)
        template_group.setLayout(template_layout)
        layout.addWidget(template_group)
        
        excel_group = QGroupBox("Excel数据文件")
        excel_layout = QVBoxLayout()
        self.excel_drop = FileDropZone()
        self.excel_drop.files_dropped.connect(self._on_excel_dropped)
        self.excel_drop.setMaximumHeight(80)
        excel_layout.addWidget(self.excel_drop)
        self.excel_label = QLabel("未选择Excel")
        self.excel_label.setStyleSheet("padding: 5px; color: #666;")
        excel_layout.addWidget(self.excel_label)
        excel_group.setLayout(excel_layout)
        layout.addWidget(excel_group)
        
        info_group = QGroupBox("使用说明")
        info_layout = QVBoxLayout()
        info_text = QTextEdit()
        info_text.setReadOnly(True)
        info_text.setMaximumHeight(80)
        info_text.setPlainText(
            "在Word模板中使用 {{列名}} 作为占位符\n"
            "例如: {{姓名}}、{{日期}}、{{金额}}\n"
            "Excel中需要有对应的列名"
        )
        info_layout.addWidget(info_text)
        info_group.setLayout(info_layout)
        layout.addWidget(info_group)
        
        gen_btn = QPushButton("批量生成文档")
        gen_btn.setStyleSheet("background-color: #0078d4; color: white; padding: 10px; font-size: 14px;")
        gen_btn.clicked.connect(self._on_generate)
        layout.addWidget(gen_btn)
        
        self.log_panel = LogPanel()
        layout.addWidget(self.log_panel)
        
        return widget
    
    def _on_template_dropped(self, filepaths):
        if filepaths:
            self.template_file = filepaths[0]
            self.template_label.setText(f"已选择: {os.path.basename(self.template_file)}")
    
    def _on_excel_dropped(self, filepaths):
        if filepaths:
            self.excel_file = filepaths[0]
            self.excel_label.setText(f"已选择: {os.path.basename(self.excel_file)}")
            try:
                if self.excel_file.endswith(".csv"):
                    self.excel_df = pd.read_csv(self.excel_file)
                else:
                    self.excel_df = pd.read_excel(self.excel_file)
                self.log_panel.info(f"已加载: {len(self.excel_df)} 行数据")
            except Exception as e:
                self.log_panel.error(f"加载失败: {str(e)}")
    
    def _on_generate(self):
        if not self.template_file:
            QMessageBox.warning(self._widget, "警告", "请先选择Word模板")
            return
        if self.excel_df is None:
            QMessageBox.warning(self._widget, "警告", "请先选择Excel数据文件")
            return
        
        output_dir = FileDialog.select_directory(self._widget, "选择输出目录")
        if not output_dir:
            return
        
        progress = ProgressDialog("生成文档中...")
        progress.show()
        
        try:
            total = len(self.excel_df)
            
            for idx, row in self.excel_df.iterrows():
                progress.set_status(f"生成 {idx + 1}/{total}")
                progress.set_progress(int((idx / total) * 100))
                
                doc = Document(self.template_file)
                
                for para in doc.paragraphs:
                    for key, value in row.items():
                        placeholder = "{{" + str(key) + "}}"
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))
                
                for table in doc.tables:
                    for table_row in table.rows:
                        for cell in table_row.cells:
                            for key, value in row.items():
                                placeholder = "{{" + str(key) + "}}"
                                if placeholder in cell.text:
                                    cell.text = cell.text.replace(placeholder, str(value))
                
                first_col = self.excel_df.columns[0]
                filename = str(row[first_col])[:30]
                output_path = os.path.join(output_dir, f"{filename}.docx")
                
                counter = 1
                while os.path.exists(output_path):
                    output_path = os.path.join(output_dir, f"{filename}_{counter}.docx")
                    counter += 1
                
                doc.save(output_path)
                progress.add_log(f"已生成: {filename}.docx")
            
            progress.complete("生成完成!")
            self.log_panel.info(f"已生成 {total} 个文档")
            QMessageBox.information(self._widget, "成功", f"已生成 {total} 个Word文档")
        except Exception as e:
            progress.close()
            self.log_panel.error(f"生成失败: {str(e)}")
            QMessageBox.critical(self._widget, "错误", f"生成失败:\n{str(e)}")
