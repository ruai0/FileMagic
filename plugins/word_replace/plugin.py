import os
import shutil
from typing import Dict, Any, List, Optional
from docx import Document
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QGroupBox, QMessageBox,
    QLineEdit, QTableWidget, QTableWidgetItem, QHeaderView
)
from PyQt6.QtCore import Qt
from ui.widgets.file_drop_zone import FileDropZone
from ui.widgets.log_panel import LogPanel
from core.plugin_manager import BasePlugin
from ui.dialogs.file_dialog import FileDialog
from ui.dialogs.progress_dialog import ProgressDialog


class Plugin(BasePlugin):
    """Word替换插件"""
    
    def __init__(self):
        super().__init__()
        self.name = "Word替换"
        self.description = "批量查找替换Word中的文本"
        self.version = "1.0.0"
        self.author = "ExcelTools"
        self._widget = None
        self._context: Optional[Dict[str, Any]] = None
        self.files: List[str] = []
    
    def initialize(self, context: Dict[str, Any]) -> None:
        super().initialize(context)
        self._context = context
    
    def get_widget(self):
        if self._widget is None:
            self._widget = self._create_widget()
        return self._widget
    
    def _create_widget(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        file_group = QGroupBox("Word文件列表")
        file_layout = QVBoxLayout()
        self.drop_zone = FileDropZone()
        self.drop_zone.files_dropped.connect(self._on_files_dropped)
        file_layout.addWidget(self.drop_zone)
        self.file_label = QLabel("0 个文件")
        self.file_label.setStyleSheet("padding: 10px; color: #666;")
        file_layout.addWidget(self.file_label)
        btn_row = QHBoxLayout()
        add_btn = QPushButton("添加文件")
        add_btn.clicked.connect(self._on_add)
        btn_row.addWidget(add_btn)
        clear_btn = QPushButton("清空")
        clear_btn.clicked.connect(lambda: self.files.clear() or self._update_label())
        btn_row.addWidget(clear_btn)
        file_layout.addLayout(btn_row)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        replace_group = QGroupBox("替换规则（可添加多条）")
        replace_layout = QVBoxLayout()
        
        self.replace_table = QTableWidget()
        self.replace_table.setColumnCount(2)
        self.replace_table.setHorizontalHeaderLabels(["查找", "替换为"])
        self.replace_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.replace_table.setMaximumHeight(120)
        replace_layout.addWidget(self.replace_table)
        
        btn_row = QHBoxLayout()
        add_rule = QPushButton("添加规则")
        add_rule.clicked.connect(self._add_rule)
        btn_row.addWidget(add_rule)
        del_rule = QPushButton("删除选中")
        del_rule.clicked.connect(self._del_rule)
        btn_row.addWidget(del_rule)
        replace_layout.addLayout(btn_row)
        
        replace_group.setLayout(replace_layout)
        layout.addWidget(replace_group)
        
        exec_btn = QPushButton("开始替换")
        exec_btn.setStyleSheet("background-color: #0078d4; color: white; padding: 10px; font-size: 14px;")
        exec_btn.clicked.connect(self._on_replace)
        layout.addWidget(exec_btn)
        
        self.log_panel = LogPanel()
        layout.addWidget(self.log_panel)
        
        self._add_rule()
        
        return widget
    
    def _on_files_dropped(self, filepaths):
        for fp in filepaths:
            if fp.lower().endswith((".docx", ".doc")) and fp not in self.files:
                self.files.append(fp)
        self._update_label()
    
    def _on_add(self):
        filepaths = FileDialog.open_files(self._widget, "选择Word文件", "Word文件 (*.docx)")
        if filepaths:
            for fp in filepaths:
                if fp not in self.files:
                    self.files.append(fp)
            self._update_label()
    
    def _update_label(self):
        self.file_label.setText(f"{len(self.files)} 个文件")
    
    def _add_rule(self):
        row = self.replace_table.rowCount()
        self.replace_table.insertRow(row)
    
    def _del_rule(self):
        row = self.replace_table.currentRow()
        if row >= 0:
            self.replace_table.removeRow(row)
    
    def _get_rules(self):
        rules = []
        for row in range(self.replace_table.rowCount()):
            find_item = self.replace_table.item(row, 0)
            replace_item = self.replace_table.item(row, 1)
            if find_item and find_item.text():
                find_text = find_item.text()
                replace_text = replace_item.text() if replace_item else ""
                rules.append((find_text, replace_text))
        return rules
    
    def _replace_in_paragraph(self, paragraph, rules):
        count = 0
        for run in paragraph.runs:
            for find_text, replace_text in rules:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    count += 1
        return count
    
    def _on_replace(self):
        if not self.files:
            QMessageBox.warning(self._widget, "警告", "请先添加Word文件")
            return
        
        rules = self._get_rules()
        if not rules:
            QMessageBox.warning(self._widget, "警告", "请添加替换规则")
            return
        
        output_dir = FileDialog.select_directory(self._widget, "选择输出目录")
        if not output_dir:
            return
        
        progress = ProgressDialog("替换中...")
        progress.show()
        
        total_replace = 0
        
        try:
            for idx, fp in enumerate(self.files):
                progress.set_status(f"处理 {idx+1}/{len(self.files)}: {os.path.basename(fp)}")
                progress.set_progress(int((idx / len(self.files)) * 100))
                
                doc = Document(fp)
                count = 0
                
                for para in doc.paragraphs:
                    count += self._replace_in_paragraph(para, rules)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                count += self._replace_in_paragraph(para, rules)
                
                output_path = os.path.join(output_dir, os.path.basename(fp))
                doc.save(output_path)
                
                total_replace += count
                self.log_panel.info(f"已处理: {os.path.basename(fp)} ({count} 处替换)")
            
            progress.complete("替换完成!")
            QMessageBox.information(self._widget, "成功", f"已处理 {len(self.files)} 个文件，共 {total_replace} 处替换")
        except Exception as e:
            progress.close()
            self.log_panel.error(f"替换失败: {str(e)}")
            QMessageBox.critical(self._widget, "错误", f"替换失败:\n{str(e)}")
