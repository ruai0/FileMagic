import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List, Optional
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QGroupBox, QMessageBox,
    QLineEdit, QSpinBox, QTableWidget, QTableWidgetItem, QHeaderView
)
from PyQt6.QtCore import Qt
from ui.widgets.file_drop_zone import FileDropZone
from ui.widgets.log_panel import LogPanel
from core.plugin_manager import BasePlugin
from ui.dialogs.file_dialog import FileDialog
from ui.dialogs.progress_dialog import ProgressDialog


class Plugin(BasePlugin):
    """Word批量水印插件"""
    
    def __init__(self):
        super().__init__()
        self.name = "Word批量水印"
        self.description = "给多个Word文档批量添加水印"
        self.version = "1.0.0"
        self.author = "ExcelTools"
        self._widget = None
        self._context = None
        self.files: List[str] = []
    
    def initialize(self, context):
        super().initialize(context)
        self._context = context
    
    def get_widget(self):
        if self._widget is None:
            self._widget = self._create_widget()
        return self._widget
    
    def _create_widget(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        file_group = QGroupBox("Word文件列表")
        file_layout = QVBoxLayout()
        self.drop_zone = FileDropZone()
        self.drop_zone.files_dropped.connect(self._on_files_dropped)
        file_layout.addWidget(self.drop_zone)
        
        self.file_table = QTableWidget()
        self.file_table.setColumnCount(1)
        self.file_table.setHorizontalHeaderLabels(["文件名"])
        self.file_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        file_layout.addWidget(self.file_table)
        
        btn_row = QHBoxLayout()
        add_btn = QPushButton("添加文件")
        add_btn.clicked.connect(self._on_add)
        btn_row.addWidget(add_btn)
        clear_btn = QPushButton("清空")
        clear_btn.clicked.connect(self._on_clear)
        btn_row.addWidget(clear_btn)
        file_layout.addLayout(btn_row)
        
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        opts_group = QGroupBox("水印设置")
        opts_layout = QVBoxLayout()
        
        text_layout = QHBoxLayout()
        text_layout.addWidget(QLabel("水印文字:"))
        self.text_input = QLineEdit()
        self.text_input.setPlaceholderText("输入水印文字")
        self.text_input.setText("机密文件")
        text_layout.addWidget(self.text_input)
        opts_layout.addLayout(text_layout)
        
        size_layout = QHBoxLayout()
        size_layout.addWidget(QLabel("字号:"))
        self.size_spin = QSpinBox()
        self.size_spin.setRange(10, 72)
        self.size_spin.setValue(40)
        size_layout.addWidget(self.size_spin)
        opts_layout.addLayout(size_layout)
        
        opts_group.setLayout(opts_layout)
        layout.addWidget(opts_group)
        
        exec_btn = QPushButton("批量添加水印")
        exec_btn.setStyleSheet("background-color: #0078d4; color: white; padding: 10px; font-size: 14px;")
        exec_btn.clicked.connect(self._on_add_watermark)
        layout.addWidget(exec_btn)
        
        self.log_panel = LogPanel()
        layout.addWidget(self.log_panel)
        
        return widget
    
    def _on_files_dropped(self, filepaths):
        for fp in filepaths:
            if fp.lower().endswith((".docx",)) and fp not in self.files:
                self.files.append(fp)
        self._update_table()
    
    def _on_add(self):
        filepaths = FileDialog.open_files(self._widget, "选择Word文件", "Word文件 (*.docx)")
        if filepaths:
            for fp in filepaths:
                if fp not in self.files:
                    self.files.append(fp)
            self._update_table()
    
    def _update_table(self):
        self.file_table.setRowCount(len(self.files))
        for row, fp in enumerate(self.files):
            self.file_table.setItem(row, 0, QTableWidgetItem(os.path.basename(fp)))
    
    def _on_clear(self):
        self.files.clear()
        self.file_table.setRowCount(0)
    
    def _add_watermark_to_doc(self, filepath, text, size):
        doc = Document(filepath)
        
        for section in doc.sections:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            
            run = paragraph.add_run(text)
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(192, 192, 192)
            
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(filepath)
    
    def _on_add_watermark(self):
        if not self.files:
            QMessageBox.warning(self._widget, "警告", "请先添加Word文件")
            return
        
        text = self.text_input.text()
        if not text:
            QMessageBox.warning(self._widget, "警告", "请输入水印文字")
            return
        
        size = self.size_spin.value()
        
        reply = QMessageBox.question(self._widget, "确认", 
            f"确定给 {len(self.files)} 个文件添加水印吗？\n（原文件将被修改）",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        progress = ProgressDialog("添加水印中...")
        progress.show()
        
        try:
            for idx, fp in enumerate(self.files):
                progress.set_status(f"处理 {idx + 1}/{len(self.files)}: {os.path.basename(fp)}")
                progress.set_progress(int((idx / len(self.files)) * 100))
                
                self._add_watermark_to_doc(fp, text, size)
                progress.add_log(f"已完成: {os.path.basename(fp)}")
            
            progress.complete("水印添加完成!")
            self.log_panel.info(f"已给 {len(self.files)} 个文件添加水印")
            QMessageBox.information(self._widget, "成功", f"已给 {len(self.files)} 个文件添加水印")
        except Exception as e:
            progress.close()
            self.log_panel.error(f"添加水印失败: {str(e)}")
            QMessageBox.critical(self._widget, "错误", f"添加水印失败:\n{str(e)}")
