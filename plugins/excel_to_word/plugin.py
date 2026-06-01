import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any, Optional
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QComboBox, QGroupBox, QMessageBox,
    QLineEdit, QCheckBox, QTextEdit
)
from PyQt6.QtCore import Qt
from ui.widgets.file_drop_zone import FileDropZone
from ui.widgets.log_panel import LogPanel
from core.plugin_manager import BasePlugin
from ui.dialogs.file_dialog import FileDialog
from ui.dialogs.progress_dialog import ProgressDialog


class Plugin(BasePlugin):
    """Excel转Word插件"""
    
    def __init__(self):
        super().__init__()
        self.name = "Excel转Word"
        self.description = "用Excel数据批量生成Word报告"
        self.version = "1.0.0"
        self.author = "ExcelTools"
        self._widget = None
        self._context = None
        self.source_file = None
        self.source_df = None
    
    def initialize(self, context):
        super().initialize(context)
        self._context = context
    
    def get_widget(self):
        if self._widget is None:
            self._widget = self._create_widget()
        return self._widget
    
    def _create_widget(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        file_group = QGroupBox("源Excel文件")
        file_layout = QVBoxLayout()
        self.drop_zone = FileDropZone()
        self.drop_zone.files_dropped.connect(self._on_file_dropped)
        file_layout.addWidget(self.drop_zone)
        self.file_label = QLabel("未选择文件")
        self.file_label.setStyleSheet("padding: 10px; color: #666;")
        file_layout.addWidget(self.file_label)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        opts_group = QGroupBox("生成选项")
        opts_layout = QVBoxLayout()
        
        title_layout = QHBoxLayout()
        title_layout.addWidget(QLabel("标题:"))
        self.title_input = QLineEdit()
        self.title_input.setPlaceholderText("输入Word文档标题")
        self.title_input.setText("数据报告")
        title_layout.addWidget(self.title_input)
        opts_layout.addLayout(title_layout)
        
        self.include_table = QCheckBox("包含数据表格")
        self.include_table.setChecked(True)
        opts_layout.addWidget(self.include_table)
        
        self.include_stats = QCheckBox("包含统计信息")
        self.include_stats.setChecked(True)
        opts_layout.addWidget(self.include_stats)
        
        self.include_chart = QCheckBox("包含图表描述")
        opts_layout.addWidget(self.include_chart)
        
        opts_group.setLayout(opts_layout)
        layout.addWidget(opts_group)
        
        gen_btn = QPushButton("生成Word文档")
        gen_btn.setStyleSheet("background-color: #0078d4; color: white; padding: 10px; font-size: 14px;")
        gen_btn.clicked.connect(self._on_generate)
        layout.addWidget(gen_btn)
        
        self.log_panel = LogPanel()
        layout.addWidget(self.log_panel)
        
        return widget
    
    def _on_file_dropped(self, filepaths):
        if filepaths:
            self.source_file = filepaths[0]
            self.file_label.setText(f"已选择: {os.path.basename(self.source_file)}")
            self._load_data()
    
    def _load_data(self):
        try:
            if self.source_file.endswith(".csv"):
                self.source_df = pd.read_csv(self.source_file)
            else:
                self.source_df = pd.read_excel(self.source_file)
            self.log_panel.info(f"已加载: {len(self.source_df)} 行, {len(self.source_df.columns)} 列")
        except Exception as e:
            self.log_panel.error(f"加载失败: {str(e)}")
    
    def _on_generate(self):
        if self.source_df is None:
            QMessageBox.warning(self._widget, "警告", "请先选择Excel文件")
            return
        
        output_path = FileDialog.save_file(self._widget, "保存Word文档", "Word文件 (*.docx)", ".docx")
        if not output_path:
            return
        
        try:
            doc = Document()
            title = self.title_input.text() or "数据报告"
            
            doc.add_heading(title, 0)
            
            doc.add_paragraph(f"数据来源: {os.path.basename(self.source_file)}")
            doc.add_paragraph(f"生成时间: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")
            
            if self.include_stats.isChecked():
                doc.add_heading("数据概览", level=1)
                doc.add_paragraph(f"总行数: {len(self.source_df)}")
                doc.add_paragraph(f"总列数: {len(self.source_df.columns)}")
                doc.add_paragraph(f"列名: {', '.join(self.source_df.columns.tolist())}")
                
                numeric_cols = self.source_df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    doc.add_heading("数值统计", level=2)
                    stats = self.source_df[numeric_cols].describe()
                    doc.add_paragraph(stats.to_string())
                doc.add_paragraph("")
            
            if self.include_table.isChecked():
                doc.add_heading("数据表格", level=1)
                
                display_df = self.source_df.head(20)
                
                table = doc.add_table(rows=len(display_df) + 1, cols=len(display_df.columns))
                table.style = 'Table Grid'
                
                for j, col in enumerate(display_df.columns):
                    cell = table.rows[0].cells[j]
                    cell.text = str(col)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for i, row in enumerate(display_df.values):
                    for j, val in enumerate(row):
                        table.rows[i + 1].cells[j].text = str(val)[:50]
                
                if len(self.source_df) > 20:
                    doc.add_paragraph(f"... 共 {len(self.source_df)} 行，仅显示前20行")
            
            doc.save(output_path)
            self.log_panel.info(f"已生成: {output_path}")
            QMessageBox.information(self._widget, "成功", f"Word文档已生成:\n{output_path}")
        except Exception as e:
            self.log_panel.error(f"生成失败: {str(e)}")
            QMessageBox.critical(self._widget, "错误", f"生成失败:\n{str(e)}")
