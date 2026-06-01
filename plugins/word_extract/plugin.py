import os
from typing import Dict, Any, Optional
from docx import Document
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QComboBox, QGroupBox, QMessageBox,
    QTabWidget, QTextEdit
)
from PyQt6.QtCore import Qt
from ui.widgets.file_drop_zone import FileDropZone
from ui.widgets.log_panel import LogPanel
from ui.widgets.data_table import DataTable
from core.plugin_manager import BasePlugin
from ui.dialogs.file_dialog import FileDialog


class Plugin(BasePlugin):
    """Word提取插件"""
    
    def __init__(self):
        super().__init__()
        self.name = "Word提取"
        self.description = "提取Word中的文字、表格、图片"
        self.version = "1.0.0"
        self.author = "ExcelTools"
        self._widget = None
        self._context: Optional[Dict[str, Any]] = None
        self.source_file = None
    
    def initialize(self, context: Dict[str, Any]) -> None:
        super().initialize(context)
        self._context = context
    
    def get_widget(self):
        if self._widget is None:
            self._widget = self._create_widget()
        return self._widget
    
    def _create_widget(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        file_group = QGroupBox("源Word文件")
        file_layout = QVBoxLayout()
        self.drop_zone = FileDropZone()
        self.drop_zone.files_dropped.connect(self._on_file_dropped)
        file_layout.addWidget(self.drop_zone)
        self.file_label = QLabel("未选择文件")
        self.file_label.setStyleSheet("padding: 10px; color: #666;")
        file_layout.addWidget(self.file_label)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        
        opts_group = QGroupBox("提取选项")
        opts_layout = QHBoxLayout()
        opts_layout.addWidget(QLabel("提取类型:"))
        self.type_combo = QComboBox()
        self.type_combo.addItems(["提取文字", "提取表格", "提取图片路径"])
        opts_layout.addWidget(self.type_combo)
        opts_group.setLayout(opts_layout)
        layout.addWidget(opts_group)
        
        extract_btn = QPushButton("开始提取")
        extract_btn.setStyleSheet("background-color: #0078d4; color: white; padding: 10px; font-size: 14px;")
        extract_btn.clicked.connect(self._on_extract)
        layout.addWidget(extract_btn)
        
        self.result_tabs = QTabWidget()
        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        self.result_tabs.addTab(self.text_edit, "提取结果")
        self.table_widget = DataTable()
        self.result_tabs.addTab(self.table_widget, "表格数据")
        layout.addWidget(self.result_tabs)
        
        self.log_panel = LogPanel()
        layout.addWidget(self.log_panel)
        
        return widget
    
    def _on_file_dropped(self, filepaths):
        if filepaths:
            self.source_file = filepaths[0]
            self.file_label.setText(f"已选择: {os.path.basename(self.source_file)}")
    
    def _on_extract(self):
        if not self.source_file:
            QMessageBox.warning(self._widget, "警告", "请先选择Word文件")
            return
        
        try:
            doc = Document(self.source_file)
            extract_type = self.type_combo.currentIndex()
            
            if extract_type == 0:
                self._extract_text(doc)
            elif extract_type == 1:
                self._extract_tables(doc)
            else:
                self._extract_images(doc)
            
            self.log_panel.info("提取完成")
        except Exception as e:
            self.log_panel.error(f"提取失败: {str(e)}")
            QMessageBox.critical(self._widget, "错误", f"提取失败:\n{str(e)}")
    
    def _extract_text(self, doc):
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        self.text_edit.setText("\n".join(text) if text else "未提取到文字")
        self.result_tabs.setCurrentIndex(0)
    
    def _extract_tables(self, doc):
        if doc.tables:
            table = doc.tables[0]
            headers = [cell.text for cell in table.rows[0].cells]
            data = []
            for row in table.rows[1:]:
                data.append([cell.text for cell in row.cells])
            
            self.table_widget.load_data(headers, data)
            self.result_tabs.setCurrentIndex(1)
            self.log_panel.info(f"提取到 {len(data)} 行表格数据")
        else:
            self.table_widget.load_data(["提示"], [["未找到表格"]])
            self.result_tabs.setCurrentIndex(1)
    
    def _extract_images(self, doc):
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                images.append(rel.target_ref)
        
        if images:
            self.text_edit.setText("\n".join(images))
        else:
            self.text_edit.setText("未找到图片")
        
        self.result_tabs.setCurrentIndex(0)
